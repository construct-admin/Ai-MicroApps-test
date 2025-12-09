# ------------------------------------------------------------------------------
# Visual Transcripts Generator (v2.1)
# ------------------------------------------------------------------------------
# Refactor date: 2025-12-09
# Refactored by: Imaad Fakier
#
# Purpose
# -------
# Streamlit entrypoint for OES' Visual Transcripts micro-app, aligned with the
# Coursera/Berkeley accessibility workflow and the behaviour of the “localhost”
# VT demo used by Marochelle.
#
# Key design goals
# ----------------
# - SRT file is REQUIRED and acts as the master timeline.
# - Users navigate the video in frame *steps* (e.g. every 50 frames).
# - Users can:
#     • Save specific frames.
#     • Crop the frame (rectangle) or use the full frame.
#     • Request GPT-4o Vision assistance per frame.
#     • Edit visual descriptions directly in-app.
# - Multiple saved frames can be used; each has:
#     • Frame index
#     • Exact timestamp (HH:MM:SS.mmm)
#     • Nearest SRT caption + its start time
#     • User-editable visual transcript text
# - Export generates a **combined audio + visual transcript (.docx)**:
#     • Full SRT timeline in order.
#     • Visual descriptions injected under the relevant captions.
#
# Security / Ops
# --------------
# - Access control via a hashed access code (ACCESS_CODE_HASH env var).
# - OpenAI API key from OPENAI_API_KEY env var.
# - VT_MODEL env var to override the default OpenAI model.
#
# Notes
# -----
# - This file follows the same style as other OES GenAI micro-apps for consistency.
# - Defaults (50 frame step, ~50 words) match Marochelle’s “50 & 50” workflow.
# ------------------------------------------------------------------------------

import os
import io
import cv2
import base64
import tempfile
import hashlib
from collections import OrderedDict
from datetime import timedelta

import numpy as np  # noqa: F401 (reserved for future image ops / cropping)
import streamlit as st
from PIL import Image
from docx import Document
from dotenv import load_dotenv
from streamlit_cropper import st_cropper

# ------------------------------------------------------------------------------
# Page setup and environment loading
# ------------------------------------------------------------------------------
st.set_page_config(page_title="VT Generator", page_icon="🖼️", layout="wide")
load_dotenv()  # Loads .env file variables into environment

APP_TITLE = "Visual Transcripts Generator"
DEFAULT_FPS_FALLBACK = 30
SUPPORTED_VIDEO_EXTS = ["mp4"]
SUPPORTED_SRT_EXTS = ["srt"]
MODEL_NAME = os.getenv("VT_MODEL", "gpt-4o")  # Allow override via .env
MAX_VIDEO_BYTES = 200 * 1024 * 1024  # 200MB upload guidance


# ------------------------------------------------------------------------------
# Authentication helpers
# ------------------------------------------------------------------------------
def sha256_hex(s: str) -> str:
    """Hash a provided access code using SHA-256. Used for secure auth comparison."""
    return hashlib.sha256(s.encode()).hexdigest()


ACCESS_CODE_HASH = os.getenv("ACCESS_CODE_HASH")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not ACCESS_CODE_HASH:
    st.error(
        "⚠️ Hashed access code not found. Please set ACCESS_CODE_HASH "
        "in your environment or Streamlit secrets."
    )
    st.stop()


# ------------------------------------------------------------------------------
# Session State initialization
# ------------------------------------------------------------------------------
def init_state() -> None:
    """Initialize session_state keys for consistent runtime behavior."""
    st.session_state.setdefault("authenticated", False)

    # Video / timing
    st.session_state.setdefault("video_path", None)
    st.session_state.setdefault("fps", DEFAULT_FPS_FALLBACK)
    st.session_state.setdefault("frame_count", 0)
    st.session_state.setdefault("frame_step", 50)  # navigation step (frames)
    st.session_state.setdefault("frame_index", 0)  # index in step units
    st.session_state.setdefault("video_ready", False)

    # SRT + subtitles
    st.session_state.setdefault("subtitles", OrderedDict())

    # Saved annotations – each entry:
    # {
    #   "frame_index": int,
    #   "seconds": float,
    #   "timestamp": str,
    #   "subtitle": str,
    #   "subtitle_start": float | None,
    #   "image": PIL.Image,
    #   "visual_text": str
    # }
    st.session_state.setdefault("annotations", [])

    # GPT settings
    # Marochelle typically uses "50 & 50" (navigation + description).
    st.session_state.setdefault("vt_word_limit", 50)

    # Cropping mode flag
    # True  → rectangle cropper UI in the main panel
    # False → use full frame as-is ("freehand" / no crop box)
    st.session_state.setdefault("use_rectangle_crop", True)


init_state()


# ------------------------------------------------------------------------------
# Access control gate
# ------------------------------------------------------------------------------
if not st.session_state.authenticated:
    st.title("🔒 Access Restricted")
    with st.form("auth_form", clear_on_submit=False):
        access_code_input = st.text_input("Enter Access Code:", type="password")
        submitted = st.form_submit_button("Submit")
    if submitted:
        if sha256_hex(access_code_input) == ACCESS_CODE_HASH:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect access code. Please try again.")
    st.stop()


# ------------------------------------------------------------------------------
# Utility functions
# ------------------------------------------------------------------------------
def seconds_to_timestamp(seconds: float) -> str:
    """Convert seconds float to HH:MM:SS.mmm formatted timestamp (zero-padded).

    This matches how timestamps appear in the combined visual + audio transcript.
    """
    td = timedelta(seconds=int(seconds))
    ms = int((seconds - int(seconds)) * 1000)
    return f"{str(td)}.{ms:03d}"


def parse_srt_bytes(srt_bytes: bytes) -> OrderedDict:
    """Parse SRT file bytes into an OrderedDict of {start_seconds: caption}.

    Design:
    - Handles multi-line captions.
    - Normalizes Windows newlines & comma-based timestamps.
    - Returns entries sorted by start time (ascending).
    """
    text = srt_bytes.decode("utf-8", errors="ignore").replace("\r\n", "\n")
    blocks, block = [], []
    for line in text.split("\n"):
        line = line.strip()
        if line:
            block.append(line)
        else:
            if block:
                blocks.append(block)
                block = []
    if block:
        blocks.append(block)

    parsed = []
    for b in blocks:
        if len(b) < 2:
            continue

        timing_line, text_lines = None, []
        for i, line in enumerate(b):
            if "-->" in line:
                timing_line = line
                text_lines = b[i + 1 :]
                break

        if not timing_line:
            continue

        start_str = timing_line.split("-->")[0].strip().replace(",", ".")
        parts = start_str.split(":")
        try:
            if len(parts) == 3:
                h, m, s = parts
                start_seconds = int(h) * 3600 + int(m) * 60 + float(s)
            else:
                m, s = parts
                start_seconds = int(m) * 60 + float(s)
        except Exception:
            # Skip malformed timestamps rather than breaking the entire parse
            continue

        caption = " ".join(t.strip() for t in text_lines if t.strip())
        parsed.append((start_seconds, caption))

    parsed.sort(key=lambda x: x[0])
    return OrderedDict(parsed)


def pil_to_base64_jpg(pil_img: Image.Image) -> str:
    """Convert a PIL image to a base64-encoded JPEG string for GPT-4o Vision."""
    buf = io.BytesIO()
    pil_img.save(buf, format="JPEG", quality=95)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


# ------------------------------------------------------------------------------
# OpenAI integration (SDK v1)
# ------------------------------------------------------------------------------
_openai_client = None


def get_openai_client():
    """Initialize and memoize OpenAI client using the environment API key.

    Notes:
    - Uses the official OpenAI Python SDK (v1-style client).
    - MODEL_NAME can be overridden via VT_MODEL in the environment.
    """
    global _openai_client
    if _openai_client is None:
        if not OPENAI_API_KEY:
            st.sidebar.error("Missing OPENAI_API_KEY in environment")
            raise RuntimeError("OPENAI_API_KEY missing")
        from openai import OpenAI  # type: ignore

        _openai_client = OpenAI(api_key=OPENAI_API_KEY)
    return _openai_client


def describe_image_with_gpt(
    pil_img: Image.Image,
    base_prompt: str,
    word_limit: int,
    max_tokens: int = 300,
) -> str:
    """Generate a visual description for a frame (or cropped region) via GPT-4o Vision.

    Parameters
    ----------
    pil_img:
        The image to be described (full frame or cropped selection).
    base_prompt:
        Instructional prompt that sets context and tone for the description.
    word_limit:
        Soft cap for description length; communicated to the model in natural language.
    max_tokens:
        Hard upper bound for the completion token count.

    Returns
    -------
    str
        Model-generated description text (stripped, never None).
    """
    base64_image = pil_to_base64_jpg(pil_img)
    client = get_openai_client()

    full_prompt = (
        f"{base_prompt}\n\n"
        f"IMPORTANT: Keep your response under approximately {word_limit} words. "
        f"Use clear, concise language suitable for a screen reader."
    )

    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": full_prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"},
                    },
                ],
            }
        ],
        max_tokens=max_tokens,
    )
    content = resp.choices[0].message.content or ""
    return content.strip()


# ------------------------------------------------------------------------------
# UI: Main layout and high-level flow
# ------------------------------------------------------------------------------
st.title(APP_TITLE)
st.caption(
    "Visual Transcripts Generator aligned to Coursera/Berkeley workflows – "
    "SRT-first, frame stepping, in-app editing, and GPT-4o Vision assistance."
)

# ------------------------------------------------------------------------------
# Global settings (frame step, word limit, cropping mode)
# ------------------------------------------------------------------------------
with st.sidebar.expander("⚙️ Settings", expanded=True):
    st.write("Tune how you navigate the video and how verbose GPT responses are.")

    st.session_state.frame_step = st.number_input(
        "Frame step (how many frames to skip between positions)",
        min_value=1,
        max_value=1000,
        value=int(st.session_state.frame_step),
        step=1,
        help=(
            "This controls how many frames the slider jumps at a time. "
            "For example, if set to 50, each step advances 50 frames."
        ),
    )

    st.session_state.vt_word_limit = st.slider(
        "Approximate word limit for each visual description",
        min_value=20,
        max_value=200,
        value=int(st.session_state.vt_word_limit),
        step=10,
    )

    st.session_state.use_rectangle_crop = st.checkbox(
        "Use rectangle crop mode (recommended)",
        value=bool(st.session_state.get("use_rectangle_crop", True)),
        help=(
            "When enabled, you can drag a rectangle over the frame to focus on a region. "
            "When disabled, the full frame is used (similar to the 'freehand' mode "
            "described in the localhost VT walkthrough)."
        ),
    )

# ------------------------------------------------------------------------------
# Uploaders for video and SRT
# ------------------------------------------------------------------------------
col_u1, col_u2 = st.columns([2, 1])
with col_u1:
    video_file = st.file_uploader(
        "🎬 Upload Video File (MP4)", type=SUPPORTED_VIDEO_EXTS
    )
with col_u2:
    srt_file = st.file_uploader(
        "📝 Upload Subtitle File (SRT – required)", type=SUPPORTED_SRT_EXTS
    )

# File size guidance and validation
if video_file is not None:
    if getattr(video_file, "size", None) and video_file.size > MAX_VIDEO_BYTES:
        st.error(
            "This tool currently supports videos up to approximately 200 MB.\n\n"
            "Please compress the video first (e.g. via LT tooling, Clipchamp / "
            "Microsoft Editor, or another compressor) and then re-upload."
        )
        # Prevent further processing with this file in this run
        video_file = None
    else:
        st.caption(
            "Tip: For smoother performance and to match VT workflows, "
            "compress videos to ≤ 200 MB before uploading."
        )

# Video preview and temporary storage
if video_file is not None:
    with st.expander("▶️ Click to preview uploaded video"):
        st.video(video_file)

    # Persist this upload to a temporary file
    temp_video_path = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4").name
    with open(temp_video_path, "wb") as f:
        f.write(video_file.read())
    st.session_state.video_path = temp_video_path

# Process button (extract metadata and parse SRT)
if st.button("🚀 Process video + subtitles"):
    if video_file is None or srt_file is None:
        st.error("Please upload BOTH a video file and an SRT file before processing.")
    else:
        # Parse subtitles
        st.session_state["subtitles"] = parse_srt_bytes(srt_file.read())

        # Read basic video metadata
        cap = cv2.VideoCapture(st.session_state.video_path)
        fps = cap.get(cv2.CAP_PROP_FPS) or DEFAULT_FPS_FALLBACK
        st.session_state["fps"] = int(round(fps))
        st.session_state["frame_count"] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        cap.release()

        # Reset per-run state
        st.session_state["video_ready"] = True
        st.session_state["frame_index"] = 0
        st.session_state["annotations"] = []

        st.success(
            f"Processed video ({st.session_state.frame_count} frames @ "
            f"{st.session_state.fps} fps) and SRT subtitles."
        )

# ------------------------------------------------------------------------------
# Sidebar transcript listing (raw SRT view)
# ------------------------------------------------------------------------------
st.sidebar.subheader("📜 Raw Subtitle Timeline (SRT)")
if st.session_state["subtitles"]:
    for ts_sec, text in st.session_state["subtitles"].items():
        st.sidebar.write(f"**{seconds_to_timestamp(ts_sec)}**")
        st.sidebar.caption(text)
else:
    st.sidebar.info("Upload and process an SRT to view its timeline here.")

# ------------------------------------------------------------------------------
# Frame navigation and capture (main pane)
# ------------------------------------------------------------------------------
if st.session_state.get("video_ready", False) and st.session_state.video_path:
    total_frames = st.session_state["frame_count"]
    step = max(1, int(st.session_state["frame_step"]))
    max_step_index = max(0, (total_frames - 1) // step)

    st.markdown("### 🎞 Frame Navigation")

    col_nav_1, col_nav_2 = st.columns([3, 1])
    with col_nav_1:
        # "frame_index" is in units of 'steps', not raw frames
        safe_max = max(1, max_step_index)
        step_index = st.slider(
            "Select frame position (in steps of the configured frame step)",
            min_value=0,
            max_value=safe_max,
            value=min(st.session_state["frame_index"], safe_max),
        )
    with col_nav_2:
        st.write(f"Total frames: `{total_frames}`")
        st.write(f"Frame step: `{step}`")

    # Convert "step index" back to actual frame number
    frame_number = step_index * step
    frame_number = min(frame_number, max(0, total_frames - 1))
    st.session_state["frame_index"] = step_index

    cap = cv2.VideoCapture(st.session_state["video_path"])
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)
    ret, frame = cap.read()
    cap.release()

    current_pil_image = None

    if ret and frame is not None:
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        current_pil_image = Image.fromarray(frame_rgb)

        # Crop behaviour depends on the "use_rectangle_crop" toggle.
        if st.session_state.get("use_rectangle_crop", True):
            # Rectangle crop mode (matches the localhost rectangle behaviour)
            st.markdown("### Select region to use for visual transcript")

            cropped_img = st_cropper(
                current_pil_image,
                realtime_update=True,
                box_color="#FF0000",  # Red selection box
                aspect_ratio=None,  # Free-form rectangle
            )

            st.image(cropped_img, caption="Cropped Region", use_column_width=True)

            # Use cropped image instead of full frame
            current_pil_image = cropped_img
        else:
            # "Freehand"/full-frame mode – no rectangle UI, just use the frame as-is
            st.markdown("### Using full frame (rectangle crop disabled)")
            st.image(
                current_pil_image,
                caption="Full Frame (no rectangle crop)",
                use_column_width=True,
            )

        current_seconds = frame_number / max(st.session_state["fps"], 1)
        current_timestamp = seconds_to_timestamp(current_seconds)
        st.info(f"Timestamp: `{current_timestamp}`")

    else:
        st.warning("Could not read this frame. Try a different position.")

    # Navigation + save controls
    nav_col1, nav_col2, nav_col3 = st.columns(3)
    with nav_col1:
        if st.button("⏮ Previous step"):
            st.session_state["frame_index"] = max(0, step_index - 1)
            st.rerun()
    with nav_col2:
        if st.button("⏭ Next step"):
            st.session_state["frame_index"] = min(max_step_index, step_index + 1)
            st.rerun()
    with nav_col3:
        if current_pil_image is not None and st.button("💾 Save this frame"):
            fps = st.session_state.get("fps", DEFAULT_FPS_FALLBACK)
            seconds = frame_number / max(fps, 1.0)
            timestamp = seconds_to_timestamp(seconds)

            subtitles = st.session_state["subtitles"]
            subtitle_text = "No subtitle"
            subtitle_start = None

            if subtitles:
                keys = list(subtitles.keys())
                candidate = None
                for k in keys:
                    if k <= seconds:
                        candidate = k
                    else:
                        break
                if candidate is not None:
                    subtitle_start = candidate
                    subtitle_text = subtitles.get(candidate, "No subtitle")

            annotation = {
                "frame_index": frame_number,
                "seconds": seconds,
                "timestamp": timestamp,
                "subtitle": subtitle_text,
                "subtitle_start": subtitle_start,
                "image": current_pil_image,
                "visual_text": "",
            }

            st.session_state["annotations"].append(annotation)
            st.success(f"Saved frame {frame_number} at {timestamp} for annotation.")
            st.rerun()


# ------------------------------------------------------------------------------
# Sidebar: Saved frames + in-app editing + GPT vision assistance
# ------------------------------------------------------------------------------
with st.sidebar:
    st.subheader("🖼 Saved Frames & Visual Transcripts")

    if not st.session_state["annotations"]:
        st.info(
            "Use the main panel to navigate the video and click 'Save this frame' "
            "to start building your visual transcript."
        )
    else:
        base_prompt = (
            "You are helping create visual descriptions for a course's accessibility "
            "materials. Describe only the key visual elements and on-screen text that "
            "are important for understanding the learning content. Write in a neutral, "
            "descriptive tone suitable for screen readers."
        )

        for i, ann in enumerate(st.session_state["annotations"]):
            st.markdown("---")

            # Frame preview
            st.image(
                ann["image"],
                caption=f"Frame {ann['frame_index']} @ {ann['timestamp']}",
                use_column_width=True,
            )

            # Nearest SRT subtitle
            if ann["subtitle"] and ann["subtitle"] != "No subtitle":
                st.caption(f"**SRT**: {ann['subtitle']}")
            else:
                st.caption("_No matching subtitle for this time._")

            # Stable key for this frame's visual text area
            text_key = f"vt_text_{i}"

            if text_key not in st.session_state:
                st.session_state[text_key] = ann.get("visual_text", "")

            st.write("Visual transcript (editable):")
            st.text_area(
                label="",
                key=text_key,
                height=120,
            )

            # Keep annotation in sync with widget state
            ann["visual_text"] = st.session_state[text_key]

            # Per-frame actions: GPT Assist + Remove
            btn_cols = st.columns(2)

            with btn_cols[0]:
                if st.button(f"✨ GPT assist #{i+1}", key=f"gpt_btn_{i}"):
                    try:
                        with st.spinner("Calling GPT-4o Vision…"):
                            response = describe_image_with_gpt(
                                ann["image"],
                                base_prompt=base_prompt,
                                word_limit=int(st.session_state["vt_word_limit"]),
                            )
                        st.session_state[text_key] = response
                        ann["visual_text"] = response
                        st.success("Updated from GPT-4o.")
                    except Exception as e:
                        st.error(f"Error calling GPT: {e}")

            with btn_cols[1]:
                if st.button(f"🗑 Remove #{i+1}", key=f"del_btn_{i}"):
                    del st.session_state["annotations"][i]
                    if text_key in st.session_state:
                        del st.session_state[text_key]
                    st.warning(f"Removed frame #{i+1} from annotations.")
                    st.rerun()


# ------------------------------------------------------------------------------
# Export: combined audio + visual transcript (.docx)
# ------------------------------------------------------------------------------
def build_docx_from_annotations(
    annotations: list,
    subtitles: OrderedDict,
) -> str:
    """Generate a combined audio + visual .docx transcript.

    Structure:
    ----------
    - Title and brief explanation line.
    - Full SRT timeline (all captions in order).
    - For any caption with one or more saved frames, insert the visual transcript(s)
      directly after that caption.

    This matches the workflow where:
    - Captions/SRT are already QC'ed upstream.
    - LD only needs to QC and lightly edit the visual descriptions.
    """
    doc = Document()
    doc.add_heading("Combined Visual and Audio Transcript", level=1)
    doc.add_paragraph("Timestamps are shown in hours, minutes, and seconds format.")

    if not subtitles:
        doc.add_paragraph("No SRT subtitles were loaded.")
        out_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        doc.save(out_path)
        return out_path

    # Group visual annotations by associated subtitle start time
    visuals_by_start = {}
    for ann in annotations:
        key = ann.get("subtitle_start")
        if key is None:
            continue
        visuals_by_start.setdefault(key, []).append(ann)

    # Walk through SRT in time order and insert visuals where present
    for start_sec, caption in subtitles.items():
        ts = seconds_to_timestamp(start_sec)

        # Audio caption line
        p = doc.add_paragraph()
        p.add_run(f"[{ts}] ").bold = True
        p.add_run(caption)

        # Any visual transcripts aligned to this caption
        if start_sec in visuals_by_start:
            # Sort visuals by the exact second of the captured frame
            for ann in sorted(visuals_by_start[start_sec], key=lambda a: a["seconds"]):
                visual_text = (ann.get("visual_text") or "").strip()
                if not visual_text:
                    continue

                vp = doc.add_paragraph()
                vp.add_run("Visual: ").bold = True
                vp.add_run(visual_text)

    out_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(out_path)
    return out_path


st.sidebar.subheader("📥 Download")
if st.sidebar.button("Generate .docx transcript"):
    path = build_docx_from_annotations(
        st.session_state["annotations"],
        st.session_state["subtitles"],
    )
    with open(path, "rb") as fh:
        st.sidebar.download_button(
            "Download combined visual + audio transcript (.docx)",
            data=fh,
            file_name="combined_visual_audio_transcript.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

# ------------------------------------------------------------------------------
# Logout control
# ------------------------------------------------------------------------------
st.sidebar.markdown("---")
st.sidebar.button(
    "Logout",
    on_click=lambda: st.session_state.update({"authenticated": False}),
    key="logout_button",
)
