# ------------------------------------------------------------------------------
# File: parsers.py
# Refactor date: 2025-11-13
# Refactored by: Imaad Fakier
#
# Purpose:
#     Provide text and DOCX parsing utilities used by the Canvas Import
#     micro-application for extracting <canvas_page> blocks authored inside
#     DOCX/PDF/HTML files. These blocks define page boundaries for automated
#     creation of Canvas LMS pages.
#
# Supported tag format:
#
#       <canvas_page>
#           ... HTML content ...
#       </canvas_page>
#
# Behaviour:
#     - Logic preserved 100% from original version.
#     - Regex is intentionally tolerant:
#           • Case-insensitive
#           • Allows attributes inside <canvas_page ...>
#           • DOTALL mode for multi-line content
#
# External dependencies:
#     - python-docx (Document)
# ------------------------------------------------------------------------------

import re
from typing import List
from docx import Document


# ==============================================================================
# Regular Expression for Canvas Page Blocks
# ==============================================================================

# Matches:
#   <canvas_page ...> ... </canvas_page>
#
# Notes:
#   - \b allows <canvas_page> or <canvas_page attr="...">
#   - DOTALL ensures multi-line matching
#   - Case-insensitive for user flexibility
#
_CANVAS_PAGE_RE = re.compile(
    r"<canvas_page\b[^>]*>(.*?)</canvas_page\s*>",
    re.IGNORECASE | re.DOTALL,
)


# ==============================================================================
# Text-based Extraction
# ==============================================================================


def extract_canvas_pages_from_text(text: str) -> List[str]:
    """
    Extract <canvas_page>...</canvas_page> blocks directly from raw text.

    Parameters:
        text (str):
            Arbitrary text content (from DOCX, PDF-to-text, HTML extraction, etc.)

    Returns:
        List[str]:
            A list of *raw* <canvas_page>...</canvas_page> blocks, with:
                - inner content preserved
                - surrounding page tags re-added
                - whitespace trimmed

    Behaviour:
        - Returns [] if text is empty/None.
        - Does not transform or sanitize HTML inside tags.
        - Downstream tools expect the tag wrappers to remain intact.
    """
    if not text:
        return []

    pages: List[str] = []

    for m in _CANVAS_PAGE_RE.finditer(text):
        inner = m.group(1).strip()
        pages.append(f"<canvas_page>\n{inner}\n</canvas_page>")

    return pages


# ==============================================================================
# DOCX-based Extraction
# ==============================================================================


def extract_canvas_pages(docx_like) -> List[str]:
    """
    Extract <canvas_page> blocks from a DOCX file.

    Parameters:
        docx_like:
            A file-like object or path acceptable to python-docx's Document().

    Returns:
        List[str]:
            Same output as extract_canvas_pages_from_text().

    Behaviour:
        - Reads all paragraphs in order.
        - Joins them with newline separators.
        - Passes the resulting text into the text-based extractor.
    """
    doc = Document(docx_like)
    text = "\n".join(p.text for p in doc.paragraphs)
    return extract_canvas_pages_from_text(text)


# ==============================================================================
# Diagnostics / Debug Helpers
# ==============================================================================


def scan_canvas_page_tags(text: str):
    """
    Count <canvas_page> start/end tags in a long text input.

    Useful for debugging missing closing tags or malformed documents.

    Returns:
        dict:
            {
                "starts": <int>,
                "ends": <int>,
                "balanced": <bool>
            }
    """
    starts = len(list(re.finditer(r"(?i)<canvas_page\b", text)))
    ends = len(list(re.finditer(r"(?i)</canvas_page\s*>", text)))
    return {"starts": starts, "ends": ends, "balanced": (starts == ends)}
