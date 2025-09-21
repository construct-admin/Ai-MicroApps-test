# app.py
import re
import json
from io import BytesIO

import streamlit as st

from utils import extract_tag

# KB helpers (make sure kb.py is the fixed version I sent)
from kb import ensure_client, create_vector_store, upload_file_to_vs, vector_store_supported
from openai import __version__ as openai_version  # for diagnostics display


# Minimal gdoc helpers (only for optional DOCX export from a GDoc URL)
from gdoc_utils import gdoc_id_from_url, fetch_docx_from_gdoc

# Module-tag parsing
from module_tags import split_text_by_module_tags

# Your parsers
from parsers import (
    extract_canvas_pages_from_text,  # used for tag module content
    extract_canvas_pages,
    scan_canvas_page_tags,            # kept as legacy fallback
)

# Canvas
from canvas_api import (
    list_modules, list_module_items, get_page_body, get_discussion_body,
    get_quiz_description, get_assignment_description, get_or_create_module,
    add_page, add_assignment, add_discussion, add_to_module
)

# Quizzes
from quizzes_classic import add_quiz, add_quiz_question
from quizzes_new import add_new_quiz, add_item_for_question  # <— new dispatcher



st.set_page_config(page_title="📄 DOCX → GPT (KB / Course Templates) → Canvas", layout="wide")
st.title("📄 Upload DOCX → Convert via GPT (KB / Course Templates) → Upload to Canvas")


def _init_state():
    defaults = {
        # Parsed + results
        "pages": [],
        "gpt_results": {},
        "visualized": False,

        # KB
        "vector_store_id": None,

        # Canvas caches
        "course_modules": [],
        "module_pages_cache": {}, "module_discussions_cache": {},
        "module_quizzes_cache": {}, "module_assignments_cache": {},
        "per_item_course_template_html": {},

        # Upload selection
        "upload_selected": set(),

        # Module-tag flow
        "tag_modules": [],                 # [{"name","text"}]
        "selected_tag_module_name": None,
        "selected_tag_module_text": None,

        # Auth
        "_sa_bytes": None,
        "_openai_key": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


_init_state()


# ──────────────────────────────────────────────────────────────────────────────
# 1) Storyboard settings (module tags only)
# ──────────────────────────────────────────────────────────────────────────────
with st.expander("📝 Storyboard Settings", expanded=True):
    sb_col1, sb_col2, sb_col3 = st.columns([1.2, 1, 1])
    with sb_col1:
        uploaded_file = st.file_uploader("Storyboard (.docx)", type="docx")
    with sb_col2:
        gdoc_url = st.text_input("Storyboard Google Doc URL (optional)")
    with sb_col3:
        sa_json = st.file_uploader("Service Account JSON (for Drive read, optional)", type=["json"])
        if sa_json is not None:
            st.session_state["_sa_bytes"] = sa_json.getvalue()

    st.markdown(
        "**Use tags in your storyboard:**\n\n"
        "`<module_name>Module Three</module_name>` … (content) … `</module>`\n\n"
        "We’ll detect these blocks and let you pick which **module** to process."
    )

    scan_col, pick_col = st.columns([1, 2])

    def _read_entire_doc_as_text() -> str:
        """Return storyboard as plain text (from uploaded DOCX or export of GDoc)."""
        try:
            from docx import Document  # python-docx
        except Exception:
            st.error("python-docx is required. Add `python-docx` to requirements.txt.")
            return ""

        if uploaded_file is not None:
            doc = Document(uploaded_file)
        elif gdoc_url and st.session_state.get("_sa_bytes"):
            fid = gdoc_id_from_url(gdoc_url)
            if not fid:
                st.error("Could not parse document id from the URL.")
                return ""
            try:
                buf = fetch_docx_from_gdoc(fid, st.session_state["_sa_bytes"])
                doc = Document(buf)
            except Exception as e:
                st.error(f"❌ Could not fetch/read Google Doc as DOCX: {e}")
                return ""
        else:
            return ""

        return "\n".join(p.text for p in doc.paragraphs)

    with scan_col:
        if st.button("🔎 Scan for <module_name>…</module> tags", use_container_width=True):
            raw_text = _read_entire_doc_as_text()
            if not raw_text.strip():
                st.warning("No storyboard loaded yet (upload a .docx or provide a GDoc URL + SA JSON).")
            else:
                mods = split_text_by_module_tags(raw_text)
                st.session_state["tag_modules"] = mods
                st.session_state["selected_tag_module_name"] = None
                st.session_state["selected_tag_module_text"] = None
                if mods:
                    st.success(f"Found {len(mods)} module section(s).")
                else:
                    st.warning("No <module_name>…</module> blocks found.")

    with pick_col:
        mods = st.session_state.get("tag_modules", [])
        if mods:
            label = st.selectbox(
                "Pick a module section",
                ["(pick a module)"] + [m["name"] for m in mods],
                key="pick_tag_module"
            )
            if label and label != "(pick a module)":
                chosen = next(m for m in mods if m["name"] == label)
                st.session_state["selected_tag_module_name"] = chosen["name"]
                st.session_state["selected_tag_module_text"] = chosen["text"]
                st.info(f"Selected: **{chosen['name']}**  — {len(chosen['text'])} chars")


# ──────────────────────────────────────────────────────────────────────────────
# 2) Knowledge Base (optional)
# ──────────────────────────────────────────────────────────────────────────────
with st.expander("📚 Knowledge Base (Vector Store) — optional", expanded=False):
    kb1, kb2, kb3, kb4 = st.columns([1,1,1,1])
    with kb1:
        existing_vs = st.text_input("Vector Store ID", value=st.session_state.get("vector_store_id") or "")
    with kb2:
        kb_docx = st.file_uploader("Upload template DOCX", type=["docx"])
    with kb3:
        kb_gdoc_url = st.text_input("Template Google Doc URL")
    with kb4:
        st.caption("Upload either DOCX or provide a Google Doc URL + SA JSON")

    # Diagnostics so you can see what SDK you’re on
    st.caption(f"OpenAI SDK version: {openai_version}")

    # Pre-flight: do we even have VS support?
    kb_client = None
    kb_supported = False
    if st.session_state.get("_openai_key"):
        try:
            kb_client = ensure_client(st.session_state["_openai_key"])
            kb_supported = vector_store_supported(kb_client)
        except Exception as e:
            st.warning(f"OpenAI client not ready: {e}")

    if not kb_supported:
        st.info(
            "Vector Stores are not available in this environment. "
            "You can still use **course templates**. "
            "To enable Knowledge Base (file_search), upgrade the SDK: "
            "`pip install --upgrade openai` and restart the app."
        )

    btns = st.columns([1,1,1])
    # Create Vector Store
    with btns[0]:
        if st.button("Create Vector Store", use_container_width=True, disabled=not (kb_client and kb_supported)):
            try:
                vs_id = create_vector_store(kb_client)
                st.session_state.vector_store_id = vs_id
                st.success(f"✅ Created Vector Store: {vs_id}")
            except Exception as e:
                st.error(f"Could not create Vector Store: {e}")

    # Upload template to VS
    with btns[1]:
        if st.button("Upload Template to KB", use_container_width=True,
                     disabled=not ((st.session_state.get("vector_store_id") or existing_vs) and kb_client and kb_supported)):
            try:
                vs_id = (st.session_state.get("vector_store_id") or existing_vs).strip()
                got = None
                if kb_docx is not None:
                    got = (BytesIO(kb_docx.getvalue()), kb_docx.name)
                elif kb_gdoc_url and st.session_state.get("_sa_bytes"):
                    fid = gdoc_id_from_url(kb_gdoc_url)
                    if fid:
                        data = fetch_docx_from_gdoc(fid, st.session_state["_sa_bytes"])
                        got = (data, "template_from_gdoc.docx")
                if not vs_id:
                    st.error("Vector Store ID missing.")
                elif not got:
                    st.error("Provide a template .docx or Google Doc URL + SA JSON.")
                else:
                    data, fname = got
                    res = upload_file_to_vs(kb_client, vs_id, data, fname)
                    status, via = res.get("status"), res.get("via","?")
                    if status == "completed":
                        st.success(f"✅ Template uploaded ({via}).")
                    elif status == "uploaded_file_only_no_vector_store_support":
                        st.warning(
                            "File uploaded to OpenAI, but Vector Stores aren’t supported in this SDK.\n"
                            "Please upgrade: `pip install --upgrade openai`.\n"
                            f"File id: {res.get('file_id')}"
                        )
                    else:
                        st.error(f"Upload error ({via}): {res.get('error','unknown')}")
            except Exception as e:
                st.error(f"Upload failed: {e}")

    # Use existing VS
    with btns[2]:
        if st.button("Use Existing VS ID", use_container_width=True):
            if existing_vs.strip():
                st.session_state.vector_store_id = existing_vs.strip()
                st.success(f"✅ Using Vector Store: {st.session_state.vector_store_id}")
            else:
                st.error("Paste a Vector Store ID first.")


# ──────────────────────────────────────────────────────────────────────────────
# 3) Canvas credentials & course structure
# ──────────────────────────────────────────────────────────────────────────────
with st.expander("🎓 Canvas Credentials & Course Structure", expanded=True):
    can1, can2, can3 = st.columns([1.2, 1, 1])
    with can1:
        canvas_domain = st.text_input("Canvas Base URL", placeholder="youruni.instructure.com")
        course_id = st.text_input("Canvas Course ID")
    with can2:
        canvas_token = st.text_input("Canvas API Token", type="password")
    with can3:
        st.write("")
    if st.button("Load Modules", use_container_width=True, disabled=not (canvas_domain and course_id and canvas_token)):
        try:
            mods = list_modules(canvas_domain, course_id, canvas_token)
            st.session_state.course_modules = [{"id": m["id"], "name": m["name"]} for m in mods]
            st.success(f"Loaded {len(mods)} module(s) from the course.")
        except Exception as e:
            st.error(f"Failed to load modules: {e}")
    if st.session_state.course_modules:
        st.caption("Existing modules:")
        st.write(", ".join([m["name"] for m in st.session_state.course_modules]))


# ──────────────────────────────────────────────────────────────────────────────
# 4) OpenAI API credentials
# ──────────────────────────────────────────────────────────────────────────────
with st.expander("🤖 OpenAI API Credentials", expanded=True):
    st.session_state["_openai_key"] = st.text_input("OpenAI API Key", type="password",
                                                    value=st.session_state.get("_openai_key", ""))


# ──────────────────────────────────────────────────────────────────────────────
# 5) Other settings
# ──────────────────────────────────────────────────────────────────────────────
with st.expander("⚙️ Other Settings", expanded=True):
    cols = st.columns([1, 1, 1])
    with cols[0]:
        use_new_quizzes = st.checkbox("Use New Quizzes (recommended)", value=True)
    with cols[1]:
        dry_run = st.checkbox("🔍 Preview only (Dry Run)", value=False)
    with cols[2]:
        st.caption("Dry run disables uploads.")


# ──────────────────────────────────────────────────────────────────────────────
# Parse storyboard (from the selected tag module)
# ──────────────────────────────────────────────────────────────────────────────
parse_cols = st.columns([1, 2])
with parse_cols[0]:
    has_story = bool(st.session_state.get("selected_tag_module_text"))
    if st.button("1️⃣ Parse storyboard", type="primary", use_container_width=True, disabled=not has_story):
        st.session_state.pages.clear()
        st.session_state.gpt_results.clear()
        st.session_state.visualized = False
        st.session_state.per_item_course_template_html.clear()
        st.session_state.upload_selected.clear()

        tag_text = st.session_state.get("selected_tag_module_text")
        tag_name = st.session_state.get("selected_tag_module_name")

        if tag_text:
            diag = scan_canvas_page_tags(tag_text)
            st.caption(f"Canvas-page tags → start: {diag['starts']}  end: {diag['ends']}  balanced: {diag['balanced']}")
        raw_pages = extract_canvas_pages_from_text(tag_text) if tag_text else []
        if not raw_pages:
            st.warning("No <canvas_page> blocks found in this module. Tags are case-insensitive. Example:\n"
                    "<canvas_page> ... </canvas_page>")
        raw_pages = extract_canvas_pages_from_text(tag_text) if tag_text else []

        # Build items with default module = selected module name
        last_known_module = tag_name or "General"
        TYPE_OPTIONS = ["page", "assignment", "discussion", "quiz"]

        for idx, block in enumerate(raw_pages):
            # robust normalization (prevents ValueError later)
            raw_page_type = extract_tag("page_type", block)
            page_type = (raw_page_type or "page").strip().lower()
            if page_type not in TYPE_OPTIONS:
                page_type = "page"

            page_title = (extract_tag("page_title", block) or f"Page {idx+1}").strip()
            module_name = (extract_tag("module_name", block) or last_known_module or "General").strip()
            page_template_name = (extract_tag("page_template", block) or "").strip()
            last_known_module = module_name

            st.session_state.pages.append({
                "index": idx,
                "raw": block,
                "page_type": page_type,
                "page_title": page_title,
                "module_name": module_name,
                "page_template_from_doc": page_template_name,
                "template_source": "kb",
                "template_module_id": None,
                "template_course_item": None,
            })

        st.success(f"✅ Parsed {len(st.session_state.pages)} item(s) from '{tag_name}'.")


# Show what was parsed
if st.session_state.pages:
    st.subheader("🔍 Parsed items (from selected module)")
    summary_rows = [
        f"- **{p['page_title']}** — *{p['page_type']}* · Module: `{p['module_name']}`"
        for p in st.session_state.pages
    ]
    st.markdown("\n".join(summary_rows))

    with st.expander("Show raw blocks parsed", expanded=False):
        for p in st.session_state.pages:
            with st.expander(f"{p['page_title']}  ({p['page_type']})", expanded=False):
                st.code(p["raw"], language="markdown")


# ──────────────────────────────────────────────────────────────────────────────
# Editable metadata & template picking
# ──────────────────────────────────────────────────────────────────────────────
if st.session_state.pages:
    st.subheader("2️⃣ Review & adjust item metadata (no GPT yet)")
    module_options = ["(pick module)"] + [m["name"] for m in st.session_state.course_modules]
    TYPE_OPTIONS = ["page", "assignment", "discussion", "quiz"]

    for i, p in enumerate(st.session_state.pages):
        header = f"Item {i+1}: {p['page_title']} ({p['page_type']}) · Module: {p['module_name']}"
        with st.expander(header, expanded=False):
            c1, c2, c3 = st.columns([1.2, 1, 1])

            with c1:
                p["page_title"] = st.text_input("Title", value=p["page_title"], key=f"title_{i}")

                # SAFE selectbox (prevents ValueError)
                curr_type = (p.get("page_type") or "page").strip().lower()
                if curr_type not in TYPE_OPTIONS:
                    curr_type = "page"
                p["page_type"] = st.selectbox(
                    "Content Type",
                    options=TYPE_OPTIONS,
                    index=TYPE_OPTIONS.index(curr_type),
                    key=f"type_{i}"
                )

            with c2:
                p["module_name"] = st.text_input("Target Module Name", value=p["module_name"], key=f"module_{i}")
                if module_options and len(module_options) > 1:
                    pick_name = st.selectbox("…or pick existing module", module_options, key=f"modpick_{i}")
                    if pick_name and pick_name != "(pick module)":
                        p["module_name"] = pick_name

            with c3:
                p["template_source"] = st.selectbox(
                    "Template Source", ["kb", "course"],
                    index=["kb", "course"].index(p.get("template_source", "kb")),
                    key=f"ts_{i}"
                )

            st.caption("Storyboard block (raw)")
            st.text_area("raw", value=p["raw"], height=150, key=f"raw_{i}")

            # Course Template Picker
            if p["template_source"] == "course" and st.session_state.course_modules and canvas_domain and course_id and canvas_token:
                st.markdown("**Course Template Picker** — choose the module where template items live, then pick one.")
                tm_cols = st.columns([1, 1, 1])
                with tm_cols[0]:
                    tm_pick = st.selectbox(
                        "Template Module", ["(pick module)"] + [m["name"] for m in st.session_state.course_modules],
                        key=f"tmpl_mod_{i}"
                    )
                if tm_pick and tm_pick != "(pick module)":
                    mod_id = None
                    for m in st.session_state.course_modules:
                        if m["name"] == tm_pick:
                            mod_id = m["id"]
                            break
                    p["template_module_id"] = mod_id
                    if mod_id:
                        if mod_id not in st.session_state.module_pages_cache:
                            items = list_module_items(canvas_domain, course_id, mod_id, canvas_token)
                            st.session_state.module_pages_cache[mod_id] = [
                                {"title": it.get("title"), "page_url": it.get("page_url")}
                                for it in items if it.get("type") == "Page" and it.get("page_url")
                            ]
                            st.session_state.module_discussions_cache[mod_id] = [
                                {"title": it.get("title"), "id": it.get("content_id")}
                                for it in items if it.get("type") == "Discussion" and it.get("content_id")
                            ]
                            st.session_state.module_quizzes_cache[mod_id] = [
                                {"title": it.get("title"), "id": it.get("content_id"), "classic": True}
                                for it in items if it.get("type") == "Quiz" and it.get("content_id")
                            ]
                            st.session_state.module_assignments_cache[mod_id] = [
                                {"title": it.get("title"), "id": it.get("content_id")}
                                for it in items if it.get("type") == "Assignment" and it.get("content_id")
                            ]

                        if p["page_type"] == "page":
                            page_opts = ["(pick page)"] + [x["title"] for x in st.session_state.module_pages_cache.get(mod_id, [])]
                            with tm_cols[1]:
                                page_pick = st.selectbox("Template Page", page_opts, key=f"tmpl_page_{i}")
                            if page_pick and page_pick != "(pick page)":
                                page_url = next(
                                    (x["page_url"] for x in st.session_state.module_pages_cache[mod_id] if x["title"] == page_pick),
                                    None
                                )
                                if page_url:
                                    html, _ = get_page_body(canvas_domain, course_id, page_url, canvas_token)
                                    st.session_state.per_item_course_template_html[i] = html or ""
                                    st.success("Loaded page template HTML.")

                        elif p["page_type"] == "discussion":
                            disc_opts = ["(pick discussion)"] + [x["title"] for x in st.session_state.module_discussions_cache.get(mod_id, [])]
                            with tm_cols[1]:
                                disc_pick = st.selectbox("Template Discussion", disc_opts, key=f"tmpl_disc_{i}")
                            if disc_pick and disc_pick != "(pick discussion)":
                                did = next(
                                    (x["id"] for x in st.session_state.module_discussions_cache[mod_id] if x["title"] == disc_pick),
                                    None
                                )
                                if did:
                                    html, _ = get_discussion_body(canvas_domain, course_id, did, canvas_token)
                                    st.session_state.per_item_course_template_html[i] = html or ""
                                    st.success("Loaded discussion template HTML.")

                        elif p["page_type"] == "quiz":
                            q_opts = ["(pick classic quiz)"] + [x["title"] for x in st.session_state.module_quizzes_cache.get(mod_id, [])]
                            a_opts = ["(pick assignment)"] + [x["title"] for x in st.session_state.module_assignments_cache.get(mod_id, [])]
                            with tm_cols[1]:
                                quiz_pick = st.selectbox("Template (Classic Quiz)", q_opts, key=f"tmpl_quiz_{i}")
                            with tm_cols[2]:
                                asg_pick = st.selectbox("Template (Assignment / New Quiz)", a_opts, key=f"tmpl_asg_{i}")
                            if quiz_pick and quiz_pick != "(pick classic quiz)":
                                qid = next(
                                    (x["id"] for x in st.session_state.module_quizzes_cache[mod_id] if x["title"] == quiz_pick),
                                    None
                                )
                                if qid:
                                    desc, _ = get_quiz_description(canvas_domain, course_id, qid, canvas_token)
                                    st.session_state.per_item_course_template_html[i] = desc or ""
                                    st.success("Loaded classic-quiz template description.")
                            elif asg_pick and asg_pick != "(pick assignment)":
                                aid = next(
                                    (x["id"] for x in st.session_state.module_assignments_cache[mod_id] if x["title"] == asg_pick),
                                    None
                                )
                                if aid:
                                    desc, _ = get_assignment_description(canvas_domain, course_id, aid, canvas_token)
                                    st.session_state.per_item_course_template_html[i] = desc or ""
                                    st.success("Loaded assignment template description.")


# ──────────────────────────────────────────────────────────────────────────────
# Visualization (GPT)
# ──────────────────────────────────────────────────────────────────────────────
if st.session_state.pages:
    st.divider()
    st.markdown("#### 🔎 Choose items to visualize (prepare HTML / quiz JSON)")
    sel_cols = st.columns([1, 1, 2])
    with sel_cols[0]:
        if st.button("Select all for visualization"):
            for i, _ in enumerate(st.session_state.pages):
                st.session_state[f"viz_sel_{i}"] = True
    with sel_cols[1]:
        if st.button("Clear viz selection"):
            for i, _ in enumerate(st.session_state.pages):
                st.session_state[f"viz_sel_{i}"] = False

    selected_indices = []
    for i, p in enumerate(st.session_state.pages):
        default_checked = st.session_state.get(f"viz_sel_{i}", False)
        checked = st.checkbox(
            f"{p['page_title']}  ({p['page_type']}) · Module: {p['module_name']}",
            value=default_checked, key=f"viz_sel_{i}"
        )
        if checked:
            selected_indices.append(i)

    if st.button("🔎 Visualize selected (no upload)", type="primary", use_container_width=True,
                 disabled=not (st.session_state.get("_openai_key") and selected_indices)):
        client = ensure_client(st.session_state.get("_openai_key", ""))
        for idx in selected_indices:
            p = st.session_state.pages[idx]
            raw_block = p["raw"]

            base_rules = (
                "You are an expert Canvas HTML generator.\n"
                "- Preserve ALL <a href> links and any <img> or <table> in the storyboard.\n"
                "- Replace only inner content of template areas; keep structure/classes/attributes intact.\n"
                "  if a section has no content, remove the template section in place; append extra sections at the end.\n"
                "- if a section does not exist in the template, create it with the same structure.\n"
                "- <element_type> tags are used to mark template code associations found within the file_search.\n"
                "- If some content does not map, append it as it appears in the storyboard."
                "- if a section does not exist in the template, create it with the same structure.\n"
                "- <element_type> tags are used to mark template code associations found within the file_search.\n"
                "- <accordion_title> are used for the summary tag in html accordions.\n"
                "- <accordion_content> are used for the content inside the accordion.\n"
                "- table formatting must be converted to HTML tables with <table>, <tr>, <td> tags.\n"
                "- <Table with Row Striping> is a tag and there is template code for it in the template document.\n"
                "- <Table with Column Striping> is a tag and there is template code for it in the template document.\n"
                "- <video> is also a tag with template code in the document. \n" 
                "- There is a possibility of elements within elements. Please add in the code accordingly. \n" 
                "- Keep .bluePageHeader, .header, .divisionLineYellow, .landingPageFooter intact.\n\n"
                "QUIZ RULES (when <page_type> is 'quiz'):\n"
                "- Questions appear between <quiz_start> and </quiz_end>.\n"
                "- <multiple_choice> blocks use '*' prefix to mark correct choices.\n"
                "- If <shuffle> appears inside a question, set \"shuffle\": true; else false.\n"
                "- Question-level feedback tags (optional):\n"
                "  <feedback_correct>...</feedback_correct>, <feedback_incorrect>...</feedback_incorrect>, <feedback_neutral>...</feedback_neutral>\n"
                "- Per-answer feedback (optional): '(feedback: ...)' after a choice line or <feedback>A: ...</feedback>.\n"
                "RETURN:\n"
                "1) Canvas-ready HTML (no code fences) and no other comments\n"
                "2) If page_type is 'quiz', append a JSON object at the very END (no extra text) with:\n"
                "- Support these Canvas-compatible question types:\n"
                "  multiple_choice_question (single correct), multiple_answers_question (checkboxes), true_false_question, "
                "  essay_question, short_answer_question (fill-in-one-blank), fill_in_multiple_blanks_question, "
                "  matching_question, numerical_question.\n"
                "- Include per-answer feedback when available, and overall feedback via a 'feedback' object "
                "(keys: 'correct','incorrect','neutral').\n"
                "JSON SCHEMA EXAMPLES (use only fields relevant to each type; keep it MINIFIED):\n"
                '{"quiz_description":"<p>Intro...</p>","questions":['
                # multiple choice
                '{"question_type":"multiple_choice_question","question_name":"...","question_text":"<p>...</p>",'
                '"answers":[{"text":"A","is_correct":false,"feedback":"<p>...</p>"},{"text":"B","is_correct":true,"feedback":"<p>...</p>"}],'
                '"shuffle":true,"feedback":{"correct":"<p>...</p>","incorrect":"<p>...</p>","neutral":"<p>...</p>"}},'
                # multiple answers (checkboxes)
                '{"question_type":"multiple_answers_question","question_name":"...","question_text":"<p>...</p>",'
                '"answers":[{"text":"A","is_correct":true,"feedback":"<p>...</p>"},{"text":"B","is_correct":true,"feedback":"<p>...</p>"},'
                '{"text":"C","is_correct":false,"feedback":"<p>...</p>"}],'
                '"feedback":{"correct":"<p>...</p>","incorrect":"<p>...</p>"}},'
                # true/false
                '{"question_type":"true_false_question","question_name":"...","question_text":"<p>...</p>",'
                '"answers":[{"text":"True","is_correct":false,"feedback":"<p>...</p>"},{"text":"False","is_correct":true,"feedback":"<p>...</p>"}],'
                '"feedback":{"correct":"<p>...</p>","incorrect":"<p>...</p>"}},'
                # essay
                '{"question_type":"essay_question","question_name":"...","question_text":"<p>...</p>",'
                '"feedback":{"neutral":"<p>Instructor graded.</p>"}},'
                # short answer (single blank; list acceptable strings)
                '{"question_type":"short_answer_question","question_name":"...","question_text":"<p>...</p>",'
                '"answers":[{"text":"chlorophyll"},{"text":"chlorophyl"}],'
                '"feedback":{"correct":"<p>...</p>","incorrect":"<p>...</p>"}},'
                # fill in multiple blanks (use {{blank_id}} in question_text; map answers by blank_id)
                '{"question_type":"fill_in_multiple_blanks_question","question_name":"...","question_text":"<p>H{{b1}}O is {{b2}}.</p>",'
                '"answers":[{"blank_id":"b1","text":"2","feedback":"<p>...</p>"},{"blank_id":"b2","text":"water","feedback":"<p>...</p>"}]},'
                # matching
                '{"question_type":"matching_question","question_name":"...","question_text":"<p>Match:</p>",'
                '"matches":[{"prompt":"H2O","match":"water","feedback":"<p>...</p>"},{"prompt":"NaCl","match":"salt","feedback":"<p>...</p>"}]},'
                # numerical (exact or exact+tolerance)
                '{"question_type":"numerical_question","question_name":"...","question_text":"<p>Speed?</p>",'
                '"numerical_answer":{"exact":12.5,"tolerance":0.5},'
                '"feedback":{"correct":"<p>...</p>","incorrect":"<p>...</p>"}}'
                "]}\n"
                "]}\n"
                "COVERAGE (NO-DROP) RULES\n"
                "- Do not omit or summarize any substantive content from the storyboard block.\n"
                "- Every sentence/line from the storyboard (between <canvas_page>…</canvas_page>) MUST appear in the output HTML.\n"
                "- If a piece of storyboard content doesn’t clearly map to a template section, append it as it appears in the storyboard.\n"
                "- Preserve the original order of content as much as possible.\n"
                "- Never remove <img>, <table>, or any explicit HTML already present in the storyboard; include them verbatim.\n"
            )

            template_html = None
            if p["template_source"] == "course":
                template_html = st.session_state.per_item_course_template_html.get(idx)
            tools = None
            if p["template_source"] == "kb" and st.session_state.get("vector_store_id"):
                tools = [{"type": "file_search", "vector_store_ids": [st.session_state["vector_store_id"]]}]

            if template_html:
                SYSTEM = base_rules + "\nUse the TEMPLATE HTML verbatim where structure exists. Return HTML only."
                USER = f"TEMPLATE HTML:\n{template_html}\n\nSTORYBOARD PAGE BLOCK:\n{raw_block}\n"
            else:
                SYSTEM = base_rules + ("\nUse file_search to locate the best matching template if available. Return HTML only." if tools else "")
                USER = f"STORYBOARD PAGE BLOCK:\n{raw_block}\n"

            from openai import OpenAI
            kwargs = {"model": "gpt-4o", "input": [{"role": "system", "content": SYSTEM},
                                                    {"role": "user", "content": USER}]}
            if tools:
                kwargs["tools"] = tools
            response = OpenAI(api_key=st.session_state.get("_openai_key", "")).responses.create(**kwargs)

            raw_out = response.output_text or ""
            cleaned = re.sub(r"```(html|json)?", "", raw_out, flags=re.IGNORECASE).strip()

            json_match = re.search(r"({[\s\S]+})\s*$", cleaned)
            quiz_json = None
            html_result = cleaned
            if json_match and p["page_type"] == "quiz":
                try:
                    quiz_json = json.loads(json_match.group(1))
                    html_result = cleaned[:json_match.start()].strip()
                except Exception:
                    quiz_json = None

            st.session_state.gpt_results[idx] = {"html": html_result, "quiz_json": quiz_json}

        st.session_state.visualized = True
        st.success("✅ Visualization complete. Previews below.")


# ──────────────────────────────────────────────────────────────────────────────
# Preview & Upload — separate panels + upload all
# ──────────────────────────────────────────────────────────────────────────────
if st.session_state.pages and st.session_state.visualized:
    st.subheader("3️⃣ Previews (post-GPT). Choose what to upload.")
    tabs = st.tabs(["Pages", "Assignments", "Discussions", "Quizzes"])
    type_map = {0: "page", 1: "assignment", 2: "discussion", 3: "quiz"}

    global_upload_btn_cols = st.columns([1, 3])
    with global_upload_btn_cols[0]:
        do_global_upload = st.button("🚀 Upload ALL Selected (across tabs)", type="secondary", disabled=False)

    module_cache = {}

    def _upload_item(p, html_result, quiz_json):
        mid = get_or_create_module(p["module_name"], canvas_domain, course_id, canvas_token, module_cache)
        if not mid:
            st.error("Module creation failed.")
            return False

        if p["page_type"] == "page":
            page_url = add_page(canvas_domain, course_id, p["page_title"], html_result, canvas_token)
            return bool(page_url and add_to_module(canvas_domain, course_id, mid, "Page", page_url, p["page_title"], canvas_token))

        if p["page_type"] == "assignment":
            aid = add_assignment(canvas_domain, course_id, p["page_title"], html_result, canvas_token)
            return bool(aid and add_to_module(canvas_domain, course_id, mid, "Assignment", aid, p["page_title"], canvas_token))

        if p["page_type"] == "discussion":
            did = add_discussion(canvas_domain, course_id, p["page_title"], html_result, canvas_token)
            return bool(did and add_to_module(canvas_domain, course_id, mid, "Discussion", did, p["page_title"], canvas_token))

        if p["page_type"] == "quiz":
            description = html_result
            if quiz_json and isinstance(quiz_json, dict) and "quiz_description" in quiz_json:
                description = quiz_json.get("quiz_description") or html_result

            if use_new_quizzes:
                q_list = (quiz_json or {}).get("questions", []) if isinstance(quiz_json, dict) else []
                unsupported = [
                    q for q in q_list
                    if q.get("question_type") not in ("multiple_choice_question", "multiple_answers_question", "true_false_question")
                ]
                if unsupported:
                    # Fallback to classic if unsupported types present
                    qid = add_quiz(canvas_domain, course_id, p["page_title"], description, canvas_token)
                    if qid:
                        for q in q_list:
                            add_quiz_question(canvas_domain, course_id, qid, q, canvas_token)
                        return add_to_module(canvas_domain, course_id, mid, "Quiz", qid, p["page_title"], canvas_token)
                    return False
                else:
                    assignment_id, err, status, raw = add_new_quiz(
                    canvas_domain, course_id, p["page_title"], description, canvas_token
                )
                if not assignment_id:
                    st.error(f"New Quiz (LTI) create failed [{status}]. {err}")
                    return False

                # Add ALL question types via dispatcher
                q_list = (quiz_json or {}).get("questions", []) if isinstance(quiz_json, dict) else []
                for pos, q in enumerate(q_list, start=1):
                    ok, dbg = add_item_for_question(canvas_domain, course_id, assignment_id, q, canvas_token, position=pos)
                    if not ok:
                        st.warning(f"Failed to add item {pos} ({q.get('question_type')}): {dbg}")

                ok = add_to_module(canvas_domain, course_id, mid, "Assignment", assignment_id, p["page_title"], canvas_token)
                if not ok:
                    st.warning("Created New Quiz but failed to add it to the module.")
                return ok

            else:  # classic quizzes path
                qid = add_quiz(canvas_domain, course_id, p["page_title"], description, canvas_token)
                if qid:
                    q_list = (quiz_json or {}).get("questions", []) if isinstance(quiz_json, dict) else []
                    for q in q_list:
                        add_quiz_question(canvas_domain, course_id, qid, q, canvas_token)
                    return add_to_module(canvas_domain, course_id, mid, "Quiz", qid, p["page_title"], canvas_token)
                return False

        return False

    for tab_idx, tab in enumerate(tabs):
        target_type = type_map[tab_idx]
        with tab:
            items = [p for p in st.session_state.pages if p["page_type"] == target_type]
            tcols = st.columns([1, 1, 2])
            with tcols[0]:
                if st.button(f"Select all in {target_type.title()}s"):
                    for p in items:
                        st.session_state.upload_selected.add(p["index"])
            with tcols[1]:
                if st.button(f"Clear selection in {target_type.title()}s"):
                    for p in items:
                        st.session_state.upload_selected.discard(p["index"])
            with tcols[2]:
                do_tab_upload = st.button(
                    f"🚀 Upload Selected {target_type.title()}s",
                    disabled=dry_run or not (canvas_domain and course_id and canvas_token)
                )

            for p in items:
                idx = p["index"]
                meta = f"{p['page_title']}  · Module: {p['module_name']}"
                with st.expander(meta, expanded=False):
                    html_result = st.session_state.gpt_results.get(idx, {}).get("html", "")
                    quiz_json = st.session_state.gpt_results.get(idx, {}).get("quiz_json")
                    st.code(html_result or "[No HTML returned]", language="html")
                    if p["page_type"] == "quiz" and quiz_json:
                        st.json(quiz_json)

                    chosen = idx in st.session_state.upload_selected
                    new_choice = st.checkbox("Select for upload", value=chosen, key=f"upsel_{idx}")
                    if new_choice:
                        st.session_state.upload_selected.add(idx)
                    else:
                        st.session_state.upload_selected.discard(idx)

                    can_upload = (not dry_run) and (canvas_domain and course_id and canvas_token)
                    if st.button(f"Upload '{p['page_title']}'", key=f"upl_{idx}", disabled=not can_upload):
                        ok = _upload_item(p, html_result, quiz_json)
                        st.success("✅ Uploaded and added to module.") if ok else st.error("❌ Upload failed.")

            if do_tab_upload and not dry_run:
                for p in items:
                    idx = p["index"]
                    if idx in st.session_state.upload_selected:
                        html_result = st.session_state.gpt_results.get(idx, {}).get("html", "")
                        quiz_json = st.session_state.gpt_results.get(idx, {}).get("quiz_json")
                        if _upload_item(p, html_result, quiz_json):
                            st.toast(f"Uploaded: {p['page_title']}", icon="✅")

    # Global upload
    if do_global_upload and not dry_run:
        for p in st.session_state.pages:
            idx = p["index"]
            if idx in st.session_state.upload_selected:
                html_result = st.session_state.gpt_results.get(idx, {}).get("html", "")
                quiz_json = st.session_state.gpt_results.get(idx, {}).get("quiz_json")
                if _upload_item(p, html_result, quiz_json):
                    st.toast(f"Uploaded: {p['page_title']}", icon="✅")


# Helpful hints
if not st.session_state.get("selected_tag_module_text"):
    st.info("Scan for `<module_name>…</module>` tags and pick a module, then click **Parse storyboard**.", icon="📝")
elif st.session_state.get("selected_tag_module_text") and not st.session_state.pages:
    st.warning("Click **Parse storyboard** to extract the items in the chosen module.", icon="👉")
elif st.session_state.pages and not st.session_state.visualized:
    st.info("Review & adjust metadata above (and pick course templates if desired), then click **Visualize selected**.", icon="🔎")
