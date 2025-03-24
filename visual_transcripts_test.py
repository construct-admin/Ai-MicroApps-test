import streamlit as st
import cv2
import numpy as np
import os
import tempfile
import base64
import requests
from PIL import Image
from docx import Document
from datetime import timedelta
import hashlib


st.set_page_config(page_title="VT Generator", page_icon="ðŸ–¼ï¸", layout="wide")

### Hash function for access code encryption
def hash_code(input_code):
    """Hashes the access code using SHA-256."""
    return hashlib.sha256(input_code.encode()).hexdigest()

### Retrieve hash code from environment variable
ACCESS_CODE_HASH = os.getenv("ACCESS_CODE_HASH")

if not ACCESS_CODE_HASH:
    st.error("âš ï¸ Hashed access code not found. Please set ACCESS_CODE_HASH.")
    st.stop()

### Authentication Logic
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("ðŸ”’ Access Restricted")
    access_code_input = st.text_input("Enter Access Code:", type="password", key="access_code_input")

    if st.button("Submit", key="submit_access_code"):
        if hash_code(access_code_input) == ACCESS_CODE_HASH:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect access code. Please try again.")

    st.stop()  # Prevent unauthorized access

# App state
st.session_state.setdefault("saved_frames", [])
st.session_state.setdefault("saved_subtitles", [])
st.session_state.setdefault("frame_index", 0)
st.session_state.setdefault("subtitles", {})
st.session_state.setdefault("video_path", None)
st.session_state.setdefault("fps", 30)
st.session_state.setdefault("frame_count", 0)
st.session_state.setdefault("video_ready", False)

# Upload files
video_file = st.file_uploader("Upload Video File (MP4)", type=["mp4"])
srt_file = st.file_uploader("Upload Subtitle File (SRT)", type=["srt"])

# Show video in accordion
if video_file:
    with st.expander("â–¶ï¸ Click to Preview Uploaded Video"):
        st.video(video_file)
    temp_video_path = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4").name
    with open(temp_video_path, "wb") as f:
        f.write(video_file.read())
    st.session_state["video_path"] = temp_video_path

# SRT parsing
def parse_srt(file):
    subtitles = {}
    lines = file.read().decode("utf-8").split("\n")
    index, start_time = None, None
    for line in lines:
        line = line.strip()
        if line.isdigit():
            index = int(line)
        elif "-->" in line:
            start_time = line.split(" --> ")[0]
            start_time = sum(float(x) * 60 ** i for i, x in enumerate(reversed(start_time.replace(',', '.').split(':'))))
        elif line:
            if index is not None and start_time is not None:
                subtitles[start_time] = line
    return subtitles

def seconds_to_timestamp(seconds):
    td = timedelta(seconds=int(seconds))
    ms = int((seconds - int(seconds)) * 1000)
    return f"{str(td)}.{ms:03d}"

# Process button
if video_file and srt_file and st.button("Process"):
    st.session_state["subtitles"] = parse_srt(srt_file)
    cap = cv2.VideoCapture(st.session_state["video_path"])
    st.session_state["fps"] = int(cap.get(cv2.CAP_PROP_FPS))
    st.session_state["frame_count"] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    cap.release()
    st.session_state["video_ready"] = True
    st.success("Video and subtitles processed.")

# Sidebar transcript view
st.sidebar.subheader("Transcript")
for timestamp, text in st.session_state["subtitles"].items():
    st.sidebar.write(f"**{timestamp}**: {text}")

# ðŸ”’ Only show frame section after processing
if st.session_state.get("video_ready", False):

    # Frame navigation
    frame_slider = st.slider("Select Frame", 0, st.session_state["frame_count"] - 1, st.session_state["frame_index"])
    st.session_state["frame_index"] = frame_slider

    cap = cv2.VideoCapture(st.session_state["video_path"])
    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_slider)
    ret, frame = cap.read()
    cap.release()

    if ret:
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        pil_image = Image.fromarray(frame_rgb)
        st.image(pil_image, caption=f"Frame {frame_slider}")
    else:
        st.warning("Could not read frame.")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Previous Frame"):
            st.session_state["frame_index"] = max(0, st.session_state["frame_index"] - 1)
    with col2:
        if st.button("Next Frame"):
            st.session_state["frame_index"] = min(st.session_state["frame_index"] + 1, st.session_state["frame_count"] - 1)

    # Save current frame
    if st.button("Save This Frame"):
        st.session_state["saved_frames"].append({
            "image": pil_image,
            "original_frame_index": frame_slider
        })
        subtitle = next((text for time, text in st.session_state["subtitles"].items()
                        if int(time * st.session_state["fps"]) == frame_slider), "No Subtitle")
        st.session_state["saved_subtitles"].append(subtitle)

# Show saved frames
for i, (frame_data, subtitle) in enumerate(zip(st.session_state["saved_frames"], st.session_state["saved_subtitles"])):
    st.sidebar.image(frame_data["image"], caption=f"Saved Frame {i}")
    st.sidebar.write(subtitle)

# Encode image for GPT
def encode_image(image):
    buffered = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    image.save(buffered, format="JPEG")
    with open(buffered.name, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

# GPT transcription
if "transcriptions" not in st.session_state:
    st.session_state["transcriptions"] = {}

for i, frame_data in enumerate(st.session_state["saved_frames"]):
    if st.sidebar.button(f"Transcribe Frame {i}"):
        st.sidebar.write(f"Transcribing Frame {i}...")

        GPT_API_KEY = os.getenv("OPENAI_API_KEY")
        if not GPT_API_KEY:
            st.sidebar.error("Missing OPENAI_API_KEY")
            break

        base64_image = encode_image(frame_data["image"])
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {GPT_API_KEY}"
        }
        payload = {
            "model": "gpt-4o",
            "messages": [
                {"role": "user", "content": [
                    {"type": "text", "text": "Whatâ€™s in this image?"},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]}
            ],
            "max_tokens": 300
        }

        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        try:
            gpt_response = response.json()
            transcription = gpt_response['choices'][0]['message']['content']
            st.sidebar.text_area(f"GPT Response for Frame {i}", transcription)
            st.session_state["transcriptions"][i] = transcription
        except Exception as e:
            st.sidebar.error("Failed GPT response.")
            st.sidebar.write(response.text)

    if i in st.session_state["transcriptions"]:
        if st.sidebar.button(f"Insert into Transcript {i}"):
            fps = st.session_state.get("fps", 30)
            original_frame = frame_data["original_frame_index"]
            seconds = original_frame / fps
            timestamp = seconds_to_timestamp(seconds)
            gpt_text = f"[GPT]: {st.session_state['transcriptions'][i]}"
            if timestamp in st.session_state["subtitles"]:
                st.session_state["subtitles"][timestamp] += f"\n{gpt_text}"
            else:
                st.session_state["subtitles"][timestamp] = gpt_text
            st.sidebar.success(f"Inserted at {timestamp}")

# Download transcript
def download_transcript():
    doc = Document()
    doc.add_heading("Visual Transcript", level=1)
    for timestamp, text in st.session_state["subtitles"].items():
        doc.add_paragraph(f"{timestamp}: {text}")
    temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_doc_path)
    with open(temp_doc_path, "rb") as doc_file:
        st.sidebar.download_button("Download Transcript", doc_file, file_name="visual_transcript.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.sidebar.subheader("Download Options")
download_transcript()

### Logout Button in Sidebar
st.sidebar.button("Logout", on_click=lambda: st.session_state.update({"authenticated": False}), key="logout_button")