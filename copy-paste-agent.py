# ─────────────────────────────────────────────────────────────
# Refactor date: 2025-11-12
# Refactored by: Imaad Fakier
# Purpose: 📋 Coursera → Doc Agent - Automates extraction of
# Coursera content into .docx
# ─────────────────────────────────────────────────────────────

import re
import os
import time
import pyperclip
from dataclasses import dataclass, field
from typing import List, Optional

import streamlit as st
import pandas as pd
from playwright.sync_api import sync_playwright, TimeoutError as PWTimeoutError
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────
# ⚙️ Page Configuration
# ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Coursera → Doc Agent",
    page_icon="📋",
    layout="wide",
)
st.title("📋 Coursera → Google Doc (copy-paste) Agent")
st.caption(
    "Automates extracting Coursera module content and building a .docx ready for import into Google Docs."
)

# ─────────────────────────────────────────────────────────────
# 📂 Data Structures
# ─────────────────────────────────────────────────────────────
COURSE_EDIT_TMPL = "https://www.coursera.org/teach/{slug}/content/edit"


@dataclass
class ItemRecord:
    module_index: int
    module_title: str
    item_index: int
    item_type: str
    item_title: str
    content_text: str


@dataclass
class RunResult:
    course_slug: str
    items: List[ItemRecord] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────
# 🧩 Utilities
# ─────────────────────────────────────────────────────────────
def parse_slug_from_input(user_input: str) -> str:
    url = user_input.strip().strip('"')
    m = re.search(r"/teach/([\w\-]+)/content/edit", url)
    return m.group(1) if m else url.split("?")[0].split("/")[-1]


def wait_for_editor_ready(page, ms: int = 800):
    """Give the page time to stabilize and scroll into view."""
    time.sleep(ms / 1000.0)
    try:
        page.mouse.wheel(0, 300)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────
# 🧠 JavaScript Evaluation Scripts
# ─────────────────────────────────────────────────────────────
def deep_query_eval_script() -> str:
    """Traverse Shadow DOM to detect modules and items."""
    return r"""(function() {
        function allDeepChildren(root) {
          const out = [];
          function walk(node) {
            out.push(node);
            if (node.shadowRoot) Array.from(node.shadowRoot.children).forEach(walk);
            Array.from(node.children).forEach(walk);
          }
          walk(root);
          return out;
        }
        function text(el){return (el && (el.textContent||"")).trim().replace(/\s+\n/g,"\n").replace(/[ \t]+/g," ").trim();}
        const root=document.documentElement;
        const nodes=allDeepChildren(root);
        const moduleRows=nodes.filter(n=>n.getAttribute&&(
          n.getAttribute("data-e2e")==="module-row"||n.getAttribute("data-e2e")==="content-module"
        ));
        const outline=[];
        for(const m of moduleRows){
          const kids=allDeepChildren(m);
          const titleEl=kids.find(k=>k.getAttribute&&k.getAttribute("data-e2e")==="module-title")||m;
          const moduleTitle=(text(titleEl)||"Untitled Module");
          const itemNodes=kids.filter(k=>k.getAttribute&&k.getAttribute("data-e2e")==="module-item");
          const items=itemNodes.map((it,idx)=>{
            const leafKids=allDeepChildren(it);
            const t=text(leafKids.find(a=>a.tagName==="A"||a.tagName==="DIV")||it)||`Item ${idx+1}`;
            const lower=t.toLowerCase();
            let itype="page";
            if(lower.match(/quiz|knowledge\s*check|assessment/)) itype="quiz";
            else if(lower.match(/assignment/)) itype="assignment";
            else if(lower.match(/discussion/)) itype="discussion";
            else if(lower.match(/video|lecture/)) itype="video";
            return {title:t, type:itype};
          });
          outline.push({moduleTitle, items});
        }
        return outline;
    })();"""


def grab_visible_text_script() -> str:
    """Return visible, relevant text for the current item."""
    return r"""(function(){
      function allDeep(root){const out=[];function walk(n){out.push(n);if(n.shadowRoot)Array.from(n.shadowRoot.children).forEach(walk);Array.from(n.children).forEach(walk);}walk(root);return out;}
      const nodes=allDeep(document.documentElement);
      function scoreNode(n){const style=window.getComputedStyle(n);if(style&&(style.visibility==="hidden"||style.display==="none"))return 0;
        let t=(n.innerText||"").trim();if(!t)return 0;
        const tag=(n.tagName||"").toLowerCase();if(["nav","header","footer","button"].includes(tag))return 0;
        const cls=(n.className||"").toString().toLowerCase();if(cls.match(/toolbar|menu|toast|breadcrumb/))return 0;
        return t.length;}
      let best=null,bestScore=0;for(const n of nodes){const sc=scoreNode(n);if(sc>bestScore){best=n;bestScore=sc;}}
      const txt=best?(best.innerText||"").trim():"";
      return txt.replace(/\n{3,}/g,"\n\n");
    })();"""


# ─────────────────────────────────────────────────────────────
# 📝 Export Function
# ─────────────────────────────────────────────────────────────
def export_docx(run: RunResult, outfile: str) -> str:
    os.makedirs("generated_exports", exist_ok=True)
    path = os.path.join("generated_exports", outfile)

    doc = Document()
    doc.add_heading(f"Coursera Export: {run.course_slug}", 0)
    doc.add_paragraph("")

    current_module = None
    for rec in run.items:
        if rec.module_title != current_module:
            doc.add_heading(f"Module {rec.module_index}: {rec.module_title}", level=1)
            current_module = rec.module_title
        doc.add_heading(f"{rec.item_type.title()}: {rec.item_title}", level=2)
        content = (
            rec.content_text.strip()
            or "[Captured title only or non-text content (quiz/LTI/external).]"
        )
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph("")

    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────
# 🤖 Main Agent Routine
# ─────────────────────────────────────────────────────────────
def run_agent(
    course_input: str,
    chrome_profile: str,
    max_items: Optional[int] = None,
    use_copy_fallback: bool = True,
) -> RunResult:
    slug = parse_slug_from_input(course_input)
    result = RunResult(course_slug=slug)
    edit_url = COURSE_EDIT_TMPL.format(slug=slug)

    with sync_playwright() as p:
        try:
            ctx = p.chromium.launch_persistent_context(
                user_data_dir=chrome_profile,
                channel="chrome",
                headless=False,
                args=["--disable-blink-features=AutomationControlled"],
            )
        except Exception as e:
            result.errors.append(f"❌ Failed to open Chrome profile: {e}")
            return result

        page = ctx.new_page()
        try:
            page.goto(edit_url, wait_until="domcontentloaded", timeout=120_000)
        except Exception as e:
            result.errors.append(f"❌ Failed to open {edit_url}: {e}")
            ctx.close()
            return result

        time.sleep(4_000 / 1000)
        try:
            outline = page.evaluate(deep_query_eval_script())
        except Exception as e:
            result.errors.append(f"⚠️ Failed to build outline: {e}")
            ctx.close()
            return result

        count = 0
        for m_i, mod in enumerate(outline, start=1):
            module_title = mod.get("moduleTitle", f"Module {m_i}")
            items = mod.get("items", [])
            for it_i, item in enumerate(items, start=1):
                if max_items and count >= max_items:
                    break
                title = item.get("title", f"Item {it_i}")
                itype = item.get("type", "page")

                try:
                    page.get_by_text(title, exact=False).first.click(timeout=7_000)
                    clicked = True
                except PWTimeoutError:
                    result.errors.append(f"Could not click item: {title}")
                    clicked = False

                wait_for_editor_ready(page)

                content_text = ""
                if clicked:
                    try:
                        content_text = page.evaluate(grab_visible_text_script()) or ""
                    except Exception:
                        pass

                    if use_copy_fallback and len(content_text.strip()) < 40:
                        try:
                            page.keyboard.down("Control")
                            page.keyboard.press("KeyA")
                            page.keyboard.press("KeyC")
                            page.keyboard.up("Control")
                            cp = pyperclip.paste() or ""
                            if len(cp.strip()) > len(content_text.strip()):
                                content_text = cp
                        except Exception:
                            pass

                result.items.append(
                    ItemRecord(m_i, module_title, it_i, itype, title, content_text)
                )
                count += 1
            if max_items and count >= max_items:
                break

        ctx.close()
        return result


# ─────────────────────────────────────────────────────────────
# 🧭 Streamlit Front-End
# ─────────────────────────────────────────────────────────────
st.sidebar.header("Settings")
course_input = st.text_input(
    "Coursera course URL or slug", "technical-assessment-testing-sandbox"
)

# Cross-platform Chrome profile detection
default_profile = (
    os.path.expanduser("~/Library/Application Support/Google/Chrome")  # macOS
    if os.name == "posix"
    else os.path.expanduser(r"~\AppData\Local\Google\Chrome\User Data")  # Windows
)

chrome_profile = st.text_input("Chrome User Data directory", value=default_profile)
colA, colB, colC = st.columns([1, 1, 1])
with colA:
    max_items = st.number_input("Limit items (0 = all)", 0, 10_000, 0)
with colB:
    use_copy = st.checkbox("Enable copy-paste fallback", True)
with colC:
    export_name = st.text_input("Output file name", "coursera_export.docx")

run_btn = st.button("▶️ Run Agent", type="primary")

# ─────────────────────────────────────────────────────────────
# 🚀 Execution & Output
# ─────────────────────────────────────────────────────────────
if run_btn:
    if not course_input.strip():
        st.error("Please enter a course URL or slug.")
    elif not os.path.isdir(chrome_profile):
        st.error("Chrome profile path doesn't exist.")
    else:
        with st.status(
            "Launching Chrome and building outline…", expanded=True
        ) as status:
            st.write(f"**Target:** {course_input}")
            st.write(f"**Chrome profile:** `{chrome_profile}`")
            try:
                rr = run_agent(
                    course_input, chrome_profile, max_items or None, use_copy
                )
            except Exception as e:
                st.exception(e)
                st.stop()

            if rr.errors:
                st.warning("Some steps reported issues:")
                for e in rr.errors:
                    st.write("• " + e)

            st.write(f"Captured **{len(rr.items)}** items.")
            status.update(label="Exporting .docx…", state="running")

        outfile = export_docx(rr, export_name)
        with open(outfile, "rb") as f:
            st.download_button(
                "⬇️ Download .docx", f.read(), file_name=os.path.basename(outfile)
            )

        st.subheader("Preview")
        df = pd.DataFrame(
            [
                {
                    "Module #": it.module_index,
                    "Module": it.module_title,
                    "Type": it.item_type,
                    "Title": it.item_title,
                    "Snippet": (
                        it.content_text[:180]
                        + ("…" if len(it.content_text) > 180 else "")
                    ).replace("\n", " "),
                }
                for it in rr.items
            ]
        )
        st.dataframe(df, use_container_width=True)
        st.info(
            "Upload the .docx to Google Drive and open as Google Doc, or import manually."
        )

# ─────────────────────────────────────────────────────────────
# 📘 Help Section
# ─────────────────────────────────────────────────────────────
st.divider()
st.markdown(
    """
**Setup Notes**
1. Install dependencies:  
   `pip install streamlit playwright python-docx pyperclip`  
   `python -m playwright install chrome`
2. Close all Chrome windows before running.
3. Run with:  
   `streamlit run copy-paste-agent.py`

**How it works**
- Uses your real Chrome profile for Coursera access.
- Traverses modules and items via Shadow DOM.
- Captures visible text or clipboard fallback.
- Builds a structured `.docx` export.

**Limitations**
- Quizzes/LTI/external embeds may not expose text.
- Coursera UI updates may require selector tweaks.
"""
)
